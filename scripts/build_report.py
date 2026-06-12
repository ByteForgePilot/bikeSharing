"""
将13个报告章节合并为Word文档。
"""
import os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

BASE = r'd:\bicycle'
FILES = [
    ('报告_第一章_绪论.txt', '第一章  绪论'),
    ('报告_第二章_系统总体设计.txt', '第二章  系统总体设计'),
    ('报告_第三章_Android采集子系统.txt', '第三章  Android多传感器数据采集子系统'),
    ('报告_第四章_用户界面设计.txt', '第四章  用户界面设计'),
    ('报告_第五章_数据预处理.txt', '第五章  数据格式与预处理'),
    ('报告_第六章_F1轮胎偏摆检测.txt', '第六章  F1—轮胎偏摆检测'),
    ('报告_第七章_F2链条异响检测.txt', '第七章  F2—链条异响检测'),
    ('报告_第八章_F3车头不正检测.txt', '第八章  F3—车头不正检测'),
    ('报告_第九章_综合健康评分.txt', '第九章  综合健康评分机制'),
    ('报告_第十章_Web可视化仪表盘.txt', '第十章  Web可视化仪表盘'),
    ('报告_第十一章_Bug分析与工程实践.txt', '第十一章  Bug分析与工程实践'),
    ('报告_第十二章_实验验证与分析.txt', '第十二章  实验验证与分析'),
    ('报告_第十三章_总结与展望.txt', '第十三章  总结与展望'),
]

plt.rcParams['font.sans-serif'] = ['SimSun', 'STSong', 'SimHei']
plt.rcParams['axes.unicode_minus'] = False

IMG_DIR = os.path.join(BASE, 'report_images')
os.makedirs(IMG_DIR, exist_ok=True)
img_counter = [0]

def save_fig():
    img_counter[0] += 1
    p = os.path.join(IMG_DIR, f'fig_{img_counter[0]}.png')
    plt.savefig(p, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return p

# ─── 图表生成 ───

def flow_box(ax, x, y, w, h, text):
    ax.add_patch(plt.Rectangle((x - w/2, y - h/2), w, h, fill=True,
        facecolor='#F5F5F5', edgecolor='#666', linewidth=1))
    ax.text(x, y, text, ha='center', va='center', fontsize=8.5, family='SimSun')

def flow_arrow(ax, x1, y1, x2, y2, lw=1):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
               arrowprops=dict(arrowstyle='->', lw=lw, color='#555'))

def draw_system_arch():
    """图2-1 系统三层架构"""
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    for y_b, y_t, title, detail, y_label in [
        (0.3, 1.7, 'Android App  (Kotlin + Jetpack Compose + Foreground Service)',
         'AccelCollector 100Hz  |  GyroCollector 50Hz  |  GpsCollector 10Hz  |  AudioCollector 8kHz', '采集层'),
        (2.5, 4.0, 'Python 算法层  (NumPy + SciPy)',
         '窗口选取 → F1轮胎偏摆 → F2链条异响 → F3车头不正 → 综合评分', '算法层'),
        (4.3, 5.5, 'Web 可视化层  (FastAPI + ECharts)',
         '健康评分仪表盘  |  三维雷达图  |  信号波形与频谱面板', '可视化层'),
    ]:
        ax.fill_between([0.5, 9.5], y_b, y_t, facecolor='#F5F5F5', edgecolor='#666', linewidth=1)
        ax.text(5, (y_b + y_t) / 2 + 0.2, title, ha='center', va='center', fontsize=10, fontweight='bold', family='SimSun')
        ax.text(5, (y_b + y_t) / 2 - 0.25, detail, ha='center', va='center', fontsize=7.5, color='#555', family='SimSun')
        ax.text(0.15, (y_b + y_t) / 2, y_label, fontsize=9, fontweight='bold', family='SimSun', va='center', rotation=90)
    for y_f, y_t, l in [(1.7, 2.1, '文件\n读取'), (4.0, 4.3, 'HTTP')]:
        ax.annotate('', xy=(5, y_t), xytext=(5, y_f), arrowprops=dict(arrowstyle='->', lw=1.2, color='#555'))
        ax.text(7.5, (y_f + y_t) / 2, l, fontsize=7, color='#555', family='SimSun', va='center', ha='center')
    ax.fill_between([0.5, 9.5], 1.7, 2.1, facecolor='#EEEEEE', edgecolor='#999', linewidth=0.8)
    ax.text(5, 1.9, '传感器数据.txt  |  音频.pcm  |  音频_时间戳.csv', ha='center', va='center', fontsize=7.5, color='#333', family='SimSun')
    plt.title('图2-1  系统三层架构', fontsize=12, fontweight='bold', family='SimSun', pad=8)
    return save_fig()

def draw_f1_flow():
    """图6-1 F1轮胎偏摆检测流程"""
    fig, ax = plt.subplots(figsize=(8, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis('off')
    boxes = [
        (5, 9.2, 6, 'Z轴加速度采集 → 去直流 → 重采样100Hz'),
        (5, 8.0, 6, 'Butterworth带通滤波 2~40Hz'),
        (5, 6.8, 6, '平整路面判定：1s滑动窗口方差 < 0.5'),
        (2.5, 5.5, 3.5, 'flat ≥ 20%\n取平整路面片段'),
        (7.5, 5.5, 3.5, 'flat < 20%\n退而使用全部数据'),
        (5, 4.3, 6, '车轮转频估计：1.5~8Hz 范围 FFT 峰值搜索'),
        (5, 3.1, 6, 'FFT 提取 A1@f, A2@2f  →  P = A1 + 0.5×A2'),
        (5, 2.0, 6, '线性评分：P≤0.35→100, P≥1.50→0\n平整惩罚：flat<20% → 上限80分'),
        (5, 0.8, 4, '输出 F1 评分 (0~100)'),
    ]
    for x, y, w, t in boxes:
        flow_box(ax, x, y, w, 0.7, t)
    arrows = [(5, 8.8, 5, 8.4), (5, 7.6, 5, 7.2), (5, 6.4, 5, 6.0),
              (5, 6.0, 2.5, 5.85), (5, 6.0, 7.5, 5.85),
              (2.5, 5.15, 5, 4.65), (7.5, 5.15, 5, 4.65),
              (5, 3.9, 5, 3.45), (5, 2.7, 5, 2.35), (5, 1.6, 5, 1.15)]
    for a in arrows: flow_arrow(ax, *a)
    plt.title('图6-1  F1轮胎偏摆检测算法流程', fontsize=12, fontweight='bold', family='SimSun', pad=10)
    return save_fig()

def draw_f2_flow():
    """图7-1 F2包络谱-倒谱双通道检测流程"""
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 8.5); ax.axis('off')
    # 包络谱通道 (左)
    left_boxes = [
        (2.5, 7.6, 3.5, '8kHz音频 → 带通2~4kHz\n希尔伯特变换 → 包络'),
        (2.5, 6.6, 3.5, '低通滤波 0.5~10Hz\n包络谱 FFT → SNR@踏频'),
        (2.5, 5.4, 3.5, '谐波检测 (2f, 3f)\n相位一致性 (圆形方差)'),
        (2.5, 4.2, 3.5, '包络调制深度\nCV = std(envelope)/mean(envelope)'),
    ]
    for x, y, w, t in left_boxes:
        flow_box(ax, x, y, w, 0.6, t)
    # 倒谱通道 (右)
    right_boxes = [
        (7.5, 7.1, 3.5, 'FFT → 对数功率谱'),
        (7.5, 6.1, 3.5, 'IFFT → Cepstrum'),
        (7.5, 5.1, 3.5, '倒频率域 SNR\n(1/pedal_freq ±15%)'),
    ]
    for x, y, w, t in right_boxes:
        flow_box(ax, x, y, w, 0.6, t)
    # 左列箭头
    for y_t, y_b in [(7.3, 6.9), (6.3, 5.7), (5.1, 4.5)]:
        flow_arrow(ax, 2.5, y_t, 2.5, y_b)
    # 右列箭头
    for y_t, y_b in [(6.8, 6.4), (5.8, 5.4)]:
        flow_arrow(ax, 7.5, y_t, 7.5, y_b)
    # 融合
    ax.add_patch(plt.Rectangle((1, 2.5), 8, 0.8, fill=True, facecolor='#EEEEEE', edgecolor='#444', linewidth=1.5))
    ax.text(5, 2.9, '五特征加权融合：SNR×0.35 + mod_depth×0.25 + harm×0.10 + phase×0.10 + cepstrum×0.20',
            ha='center', va='center', fontsize=9, fontweight='bold', family='SimSun')
    flow_box(ax, 5, 1.5, 6, 0.7, '分段评分 → F2 评分 (0~100)')
    flow_arrow(ax, 5, 2.5, 5, 1.85, 1.5)
    # 通道标注
    ax.text(2.5, 8.2, '包络谱通道', ha='center', fontsize=10, fontweight='bold', family='SimSun')
    ax.text(7.5, 8.2, '倒谱通道', ha='center', fontsize=10, fontweight='bold', family='SimSun')
    plt.title('图7-1  F2包络谱-倒谱双通道检测流程', fontsize=12, fontweight='bold', family='SimSun', pad=10)
    return save_fig()

def draw_f3_flow():
    """图8-1 F3车头不正检测流程"""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 8); ax.axis('off')
    boxes = [
        (5, 7.4, 6, '陀螺仪 Z轴重采样50Hz → 5s滑动窗口 → 取gz方差最低30%为直行段'),
        (5, 6.2, 5, '质量门：最佳段 σ_gz < 0.3 rad/s ?'),
        (7.5, 5.1, 3, '是 → 继续检测'),
        (2.5, 5.1, 3, '否 → 返回满分 (数据不足)'),
        (5, 4.0, 6, '偏置估计：Δθ = deg(|median(gz_means)| × obs_time)'),
        (5, 2.9, 6, '符号一致性增强 (>70%同向时轻度放大)'),
        (5, 1.8, 6, '加速度计侧向晃动惩罚 (Δθ<3°且 stability_ratio>0.55时触发)'),
        (5, 0.7, 5, '评分：Δθ≤4°→100, Δθ≥10°→0, 中间线性'),
    ]
    for x, y, w, t in boxes:
        flow_box(ax, x, y, w, 0.7, t)
    arrows = [(5, 7.0, 5, 6.55), (5, 6.55, 7.5, 5.45), (5, 6.55, 2.5, 5.45),
              (7.5, 4.75, 5, 4.35), (5, 3.6, 5, 3.25), (5, 2.5, 5, 2.15), (5, 1.4, 5, 1.05)]
    for a in arrows: flow_arrow(ax, *a)
    plt.title('图8-1  F3车头不正检测算法流程 (v3.2)', fontsize=12, fontweight='bold', family='SimSun', pad=10)
    return save_fig()

def draw_scoring_flow():
    """图9-1 综合评分计算流程"""
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    for x, y, w, t in [(2.5, 5.2, 3, 'F1 轮胎偏摆\n权重 0.40'),
                        (5, 5.2, 3, 'F2 链条异响\n权重 0.30'),
                        (7.5, 5.2, 3, 'F3 车头不正\n权重 0.30')]:
        flow_box(ax, x, y, w, 0.7, t)
    flow_box(ax, 5, 3.8, 7, 0.7, '加权调和平均  H = 1 / (0.4/F1 + 0.3/F2 + 0.3/F3)')
    for x in [2.5, 5, 7.5]: flow_arrow(ax, x, 4.85, 5, 4.15)
    flow_box(ax, 5, 2.4, 4, 0.7, 'S = H × (min/100)^0.7')
    flow_arrow(ax, 5, 3.45, 5, 2.75)
    flow_box(ax, 5, 1.0, 7, 0.7, 'S ≥ 75 → 推荐骑行  |  50~74 → 谨慎使用  |  < 50 → 建议换车')
    flow_arrow(ax, 5, 2.05, 5, 1.35)
    plt.title('图9-1  综合评分计算流程', fontsize=12, fontweight='bold', family='SimSun', pad=8)
    return save_fig()

def draw_dashboard_layout():
    """图10-1 仪表盘布局示意"""
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 6); ax.axis('off')
    ax.add_patch(plt.Rectangle((0, 5.2), 9, 0.7, facecolor='#444', edgecolor='#444'))
    ax.text(4.5, 5.55, '共享单车健康快速检测系统', ha='center', va='center', color='white', fontsize=12, fontweight='bold', family='SimSun')
    ax.add_patch(plt.Rectangle((0.2, 0.2), 2.6, 4.8, facecolor='#F8F8F8', edgecolor='#999', linewidth=1))
    ax.text(1.5, 4.5, '健康评分\n仪表盘', ha='center', fontsize=10, fontweight='bold', family='SimSun')
    ax.add_patch(plt.Rectangle((0.4, 2.5), 2.2, 1.5, facecolor='white', edgecolor='#CCC'))
    ax.text(1.5, 3.0, '82.8', ha='center', fontsize=28, fontweight='bold', color='#444', family='SimSun')
    ax.text(1.5, 2.65, '谨慎使用', ha='center', fontsize=9, color='#888', family='SimSun')
    ax.text(1.5, 2.1, 'F1: 80  F2: 95  F3: 82', ha='center', fontsize=7, color='#999', family='SimSun')
    ax.add_patch(plt.Rectangle((3.1, 0.2), 5.7, 4.8, facecolor='#F8F8F8', edgecolor='#999', linewidth=1))
    ax.text(6, 4.5, '详细分析面板', ha='center', fontsize=10, fontweight='bold', family='SimSun')
    for i, (label, y) in enumerate([('F1 轮胎偏摆 — Z轴波形 + FFT频谱', 3.7),
                                      ('F2 链条异响 — 音频波形 + FFT频谱', 2.8),
                                      ('F3 车头不正 — 陀螺仪波形 + 累计偏航角', 1.9),
                                      ('数据摘要 (采样统计 + 窗口信息)', 1.0)]):
        ax.add_patch(plt.Rectangle((3.3, y - 0.3), 5.3, 0.55, facecolor='white', edgecolor='#DDD'))
        ax.text(3.5, y, label, fontsize=8, family='SimSun')
    plt.title('图10-1  Web仪表盘布局示意', fontsize=12, fontweight='bold', family='SimSun', pad=8)
    return save_fig()

def draw_penalty_curve():
    """图9-2 不同幂指数的penalty曲线"""
    fig, ax = plt.subplots(figsize=(7, 4))
    mins = np.linspace(0, 100, 201)
    for exp, label, ls in [(1.0, 'linear (^1.0)', '--'), (0.7, '^0.7 (selected)', '-'),
                            (0.6, '^0.6', ':'), (0.5, '^0.5', '-.'),(0, 'no penalty', (0,(3,1)))]:
        p = np.ones_like(mins) if exp == 0 else (mins / 100.0) ** exp
        ax.plot(mins, p, label=label, linestyle=ls, linewidth=2)
    ax.set_xlabel('min(F1, F2, F3)', fontsize=11, family='SimSun')
    ax.set_ylabel('penalty', fontsize=11, family='SimSun')
    ax.set_title('图9-2  不同幂指数的 penalty 曲线', fontsize=12, fontweight='bold', family='SimSun')
    ax.legend(fontsize=9); ax.grid(True, alpha=0.3); ax.set_xlim(0, 100); ax.set_ylim(0, 1.05)
    return save_fig()

def draw_experiment_barchart():
    """图12-1 三组数据检测结果对比"""
    fig, ax = plt.subplots(figsize=(8, 4.5))
    datasets = ['(3) Good', '(6) Medium', '(2) Low']
    f1 = [100, 80, 71.2]; f2 = [95.5, 94.8, 92.4]; f3v = [100, 81.5, 93.7]; total = [98.6, 72.2, 65.3]
    x = np.arange(len(datasets)); w = 0.18
    colors = ['#555', '#888', '#BBB', '#333']
    for i, (vals, name) in enumerate(zip([f1, f2, f3v, total], ['F1 Tire', 'F2 Chain', 'F3 Handlebar', 'Total'])):
        b = ax.bar(x + (i - 1.5) * w, vals, w, label=name, color=colors[i], edgecolor='white')
        for rect in b:
            ax.text(rect.get_x() + rect.get_width()/2., rect.get_height() + 0.5, f'{rect.get_height():.0f}', ha='center', fontsize=7)
    ax.set_ylabel('Score (0~100)', fontsize=11, family='SimSun')
    ax.set_xticks(x); ax.set_xticklabels(datasets, fontsize=11, family='SimSun')
    ax.set_ylim(0, 115); ax.legend(fontsize=8, ncol=2)
    ax.axhline(y=75, color='#999', linestyle='--', linewidth=1)
    ax.axhline(y=50, color='#CCC', linestyle='--', linewidth=1)
    ax.grid(axis='y', alpha=0.3)
    ax.set_title('图12-1  三组数据检测结果对比', fontsize=12, fontweight='bold', family='SimSun')
    return save_fig()

def draw_sensor_sampling():
    """图3-1 传感器采样时序示意"""
    fig, ax = plt.subplots(figsize=(8, 3.5))
    ax.set_xlim(0, 1.0); ax.set_ylim(0, 4)
    for name, interval, y, c in [('Accel 100Hz', 0.01, 1.5, '#444'),
                                   ('Gyro 50Hz', 0.02, 2.5, '#777'),
                                   ('GPS 10Hz', 0.1, 3.5, '#AAA')]:
        times = np.arange(0, 1.0, interval)
        ax.vlines(times, y - 0.3, y + 0.3, colors=c, linewidths=1.5)
        ax.text(1.02, y, name, fontsize=9, color=c, va='center', fontweight='bold', family='SimSun')
    ax.set_yticks([]); ax.set_xlabel('Time (seconds)', fontsize=11, family='SimSun')
    ax.set_title('图3-1  传感器采样时序示意 (前1秒)', fontsize=12, fontweight='bold', family='SimSun')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False); ax.spines['left'].set_visible(False)
    return save_fig()

# ─── 文档解析工具 ───

def is_pipe_table_line(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and '|' in s[1:-1]

def detect_pipe_table(lines, start_idx):
    i, rows = start_idx, []
    while i < len(lines):
        s = lines[i].strip()
        if not is_pipe_table_line(s): break
        if re.match(r'^[\s|:\-]+$', s): i += 1; continue
        cells = [c.strip() for c in s.split('|')[1:-1]]
        if cells: rows.append(cells)
        i += 1
    return rows if len(rows) >= 2 else None

def convert_pipe_table(doc, rows):
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell.text = ''
            if j < len(row_data):
                run = cell.paragraphs[0].add_run(row_data[j])
                run.font.size = Pt(8); run.font.name = '宋体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if i == 0: run.font.bold = True
    doc.add_paragraph()

def is_ascii_box_line(line):
    s = line.strip()
    return any(c in s for c in '┌┐└┘├┤┬┴┼│─╞╪')

def detect_ascii_table(lines, start_idx):
    i, rows = start_idx, []
    while i < len(lines):
        s = lines[i].strip()
        if not s: break
        if not (is_ascii_box_line(s) or '│' in s): break
        cells = [c.strip() for c in s.split('│')[1:-1]]
        if cells and any(c for c in cells): rows.append(cells)
        i += 1
    return rows if rows else None

def convert_ascii_table(doc, rows):
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]; cell.text = ''
            if j < len(row_data):
                run = cell.paragraphs[0].add_run(row_data[j])
                run.font.size = Pt(8); run.font.name = '宋体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if i == 0: run.font.bold = True
    doc.add_paragraph()

# ─── 主文档构建 ───

def add_heading_style(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return h

def build_document():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'; style.font.size = Pt(11); style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.first_line_indent = Pt(22)

    for lvl in [1, 2, 3]:
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = '宋体'; hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    # 封面
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('共享单车多传感器数据采集\n与健康检测系统')
    run.font.size = Pt(26); run.font.bold = True; run.font.name = '宋体'; run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('项目技术报告'); run.font.size = Pt(18); run.font.name = '宋体'; run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年6月'); run.font.size = Pt(14); run.font.name = '宋体'; run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_page_break()

    # 目录
    h = add_heading_style(doc, '目  录', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('（在Word中插入引用→目录→自动目录）')
    doc.add_page_break()

    # 逐章节
    for filename, chapter_title in FILES:
        path = os.path.join(BASE, filename)
        if not os.path.exists(path):
            print(f'  [MISS] {filename}'); continue
        print(f'  Process: {filename}')
        text = open(path, encoding='utf-8').read()
        add_heading_style(doc, chapter_title, 1)
        lines = text.split('\n'); i = 0

        while i < len(lines):
            line = lines[i].rstrip() if i < len(lines) else ''

            if not line.strip(): i += 1; continue
            if line.strip().startswith('===') and len(line.strip()) > 10: i += 1; continue
            if '完' in line and line.strip().startswith('==='): i += 1; continue

            # 小节标题
            if re.match(r'^\d+\.\d+[\s　]', line.strip()):
                add_heading_style(doc, line.strip(), 2); i += 1; continue
            if re.match(r'^\d+\.\d+\.\d+[\s　]', line.strip()):
                add_heading_style(doc, line.strip(), 3); i += 1; continue

            # 管道表格
            if is_pipe_table_line(line.strip()):
                rows = detect_pipe_table(lines, i)
                if rows: convert_pipe_table(doc, rows); i += len(rows) + 1
                else: i += 1
                continue

            # ASCII框图表格
            if is_ascii_box_line(line.strip()):
                rows = detect_ascii_table(lines, i)
                if rows: convert_ascii_table(doc, rows); i += len(rows) + 1
                else: i += 1
                continue

            # 代码块
            if line.strip().startswith('```'): i += 1; continue

            # 累积段落
            para_lines = []
            while i < len(lines):
                l = lines[i].rstrip()
                if not l.strip(): break
                if l.strip().startswith('===') and len(l.strip()) > 10: break
                if '完' in l and l.strip().startswith('==='): break
                if re.match(r'^\d+\.\d+[\s　]', l.strip()): break
                if re.match(r'^\d+\.\d+\.\d+[\s　]', l.strip()): break
                if is_pipe_table_line(l.strip()) or is_ascii_box_line(l.strip()): break
                if l.strip().startswith('```'): break
                clean = l.strip()
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
                clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
                clean = re.sub(r'^[-•]\s+', '', clean)
                para_lines.append(clean)
                i += 1

            if para_lines:
                p = doc.add_paragraph(''.join(para_lines), style='Normal')
                for run in p.runs:
                    run.font.size = Pt(11); run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 插图
        img_chapter_map = {
            '第二章': ['system_arch'],
            '第三章': ['sensor_sampling'],
            '第六章': ['f1_flow'], '第七章': ['f2_flow'],
            '第八章': ['f3_flow'], '第九章': ['scoring_flow', 'penalty_curve'],
            '第十章': ['dashboard_layout'], '第十二章': ['experiment_barchart'],
        }
        for ch_key, imgs in img_chapter_map.items():
            if ch_key in chapter_title:
                for ik in imgs:
                    ip = os.path.join(IMG_DIR, f'{ik}.png')
                    if os.path.exists(ip):
                        doc.add_picture(ip, width=Inches(5.5))

        doc.add_page_break()

    return doc

# ═══════════════════════════════════════════════════════════

print('=' * 50)
print('Generating charts...')
chart_funcs = {
    'system_arch': draw_system_arch, 'f1_flow': draw_f1_flow,
    'f2_flow': draw_f2_flow, 'f3_flow': draw_f3_flow,
    'scoring_flow': draw_scoring_flow, 'dashboard_layout': draw_dashboard_layout,
    'penalty_curve': draw_penalty_curve, 'experiment_barchart': draw_experiment_barchart,
    'sensor_sampling': draw_sensor_sampling,
}
for name, func in chart_funcs.items():
    try:
        path = func()
        print(f'  [OK] {name}')
    except Exception as e:
        print(f'  [FAIL] {name}: {e}')

print('\nBuilding Word document...')
doc = build_document()
output_path = os.path.join(BASE, '共享单车健康检测系统_技术报告.docx')
doc.save(output_path)
print(f'[OK] Saved: {output_path}')
